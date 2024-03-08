import docx
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

class create_order:

    def replace_text(file_path, target_text, replacement_text):
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if target_text in paragraph.text:
                paragraph.text = paragraph.text.replace(target_text, replacement_text)
        doc.save(file_path)

    # Укажите путь к файлу 'ккк'
    file_path = 'путь/к/файлу/ккк.docx'
    # Укажите текст, который нужно заменить
    target_text = '===='
    # Укажите текст, на который нужно заменить
    replacement_text = '333'

    #replace_text(file_path, target_text, replacement_text)

    def format_name(self, full_name):
        # Разделение полного имени на отдельные слова
        names = full_name.split()
        # Получение фамилии
        surname = names[0]
        names = names[1:3]
        # Получение инициалов имени и отчества (первых букв остальных слов)
        initials = [name[0] + '.' for name in names]
        # Соединение фамилии и инициалов в одну строку
        formatted_name = surname + ' ' + ' '.join(initials)

        return formatted_name

    def smena(self, init):
        # Разделение полного имени на отдельные слова
        names = init.split()
        temp = names[0]
        names = names[1:3]
        names.append(temp)
        formatted_name_2 = ' '.join(map(str, names))

        return formatted_name_2

    # Назначение
    def naznach(self, fio, new_dolzn, new_otdel, data, data_zayav, slovar):
        doc = docx.Document("Приказ о назначении.docx")
        for paragraph in doc.paragraphs:
            if "#####" in paragraph.text:
                paragraph.text = paragraph.text.replace("#####", fio)
        initials = self.format_name(fio)
        for paragraph in doc.paragraphs:
            if "###" in paragraph.text:
                paragraph.text = paragraph.text.replace("###", initials)
                paragraph.runs[0].font.size = Pt(14)
        smen_init = self.smena(initials)
        for paragraph in doc.paragraphs:
            if "___" in paragraph.text:
                paragraph.text = paragraph.text.replace("___", smen_init)
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc.paragraphs:
            if "##" in paragraph.text:
                paragraph.text = paragraph.text.replace("##", data)
        for paragraph in doc.paragraphs:
            if "-_-" in paragraph.text:
                paragraph.text = paragraph.text.replace("-_-", new_dolzn)
        for paragraph in doc.paragraphs:
            if "===" in paragraph.text:
                paragraph.text = paragraph.text.replace("===", new_otdel)
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc.paragraphs:
            if "=_=" in paragraph.text:
                paragraph.text = paragraph.text.replace("=_=", data_zayav)
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc.paragraphs:
            if "head_name" in paragraph.text:
                paragraph.text = paragraph.text.replace("head_name", slovar["head_dep_name"])
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc.paragraphs:
            if "head_post" in paragraph.text:
                paragraph.text = paragraph.text.replace("head_post", slovar["head_dep_post"])
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc.paragraphs:
            if "hr_post" in paragraph.text:
                paragraph.text = paragraph.text.replace("hr_post", slovar["HR_post"])
                paragraph.runs[0].font.size = Pt(10)
        for paragraph in doc.paragraphs:
            if "hr_name" in paragraph.text:
                paragraph.text = paragraph.text.replace("hr_name", slovar["HR_name"])
                paragraph.runs[0].font.size = Pt(10)

        table = doc.tables[1]
        cell = table.cell(1, 0)
        cell.text = slovar['Curator_post']
        cell = table.cell(1, 1)
        cell.text = slovar['Curator_name']

        cell = table.cell(2, 0)
        cell.text = slovar['Law_dep_post']
        cell = table.cell(2, 1)
        cell.text = slovar['Law_dep_name']

        cell = table.cell(3, 0)
        cell.text = slovar['admin_dep_post']
        cell = table.cell(3, 1)
        cell.text = slovar['admin_dep_name']

        for row in table.rows:
            for cell in row.cells:
                paragraph = cell.paragraphs[0]
                paragraph.runs[0].font.size = Pt(14)

        paragraph_text = f'О назначении {initials}'
        for paragraph in doc.paragraphs:
            if paragraph_text in paragraph.text:
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.bold = True
        paragraph_text = f'«О назначении {initials}»'
        for paragraph in doc.paragraphs:
            if paragraph_text in paragraph.text:
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.bold = False

        doc.save(f'приказы/Приказ о назначении_{datetime.now().strftime("%Y-%m-%d %H.%M.%S")}_{initials}.docx')

    # Увольнение
    def dismissal(self, fio, dolzn, otdel, data, data_zayav, slovar):
        doc_1 = docx.Document("Приказ об увольнении.docx")
        for paragraph in doc_1.paragraphs:
            if "#####" in paragraph.text:
                paragraph.text = paragraph.text.replace("#####", fio)
                paragraph.runs[0].font.size = Pt(14)
        initials = self.format_name(fio)
        for paragraph in doc_1.paragraphs:
            if "###" in paragraph.text:
                paragraph.text = paragraph.text.replace("###", initials)
                paragraph.runs[0].font.size = Pt(14)
        smen_init = self.smena(initials)
        for paragraph in doc_1.paragraphs:
            if "___" in paragraph.text:
                paragraph.text = paragraph.text.replace("___", smen_init)
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc_1.paragraphs:
            if "##" in paragraph.text:
                paragraph.text = paragraph.text.replace("##", data)
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc_1.paragraphs:
            if "-_-" in paragraph.text:
                paragraph.text = paragraph.text.replace("-_-", dolzn)
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc_1.paragraphs:
            if "===" in paragraph.text:
                paragraph.text = paragraph.text.replace("===", otdel)
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc_1.paragraphs:
            if "=_=" in paragraph.text:
                paragraph.text = paragraph.text.replace("=_=", data_zayav)
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc_1.paragraphs:
            if "head_name" in paragraph.text:
                paragraph.text = paragraph.text.replace("head_name", slovar["head_dep_name"])
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc_1.paragraphs:
            if "head_post" in paragraph.text:
                paragraph.text = paragraph.text.replace("head_post", slovar["head_dep_post"])
                paragraph.runs[0].font.size = Pt(14)
        for paragraph in doc_1.paragraphs:
            if "hr_post" in paragraph.text:
                paragraph.text = paragraph.text.replace("hr_post", slovar["HR_post"])
                paragraph.runs[0].font.size = Pt(10)
        for paragraph in doc_1.paragraphs:
            if "hr_name" in paragraph.text:
                paragraph.text = paragraph.text.replace("hr_name", slovar["HR_name"])
                paragraph.runs[0].font.size = Pt(10)

        table_1 = doc_1.tables[1]
        cell = table_1.cell(1, 0)
        cell.text = slovar['Curator_post']
        cell = table_1.cell(1, 1)
        cell.text = slovar['Curator_name']

        cell = table_1.cell(2, 0)
        cell.text = slovar['Law_dep_post']
        cell = table_1.cell(2, 1)
        cell.text = slovar['Law_dep_name']

        cell = table_1.cell(3, 0)
        cell.text = slovar['admin_dep_post']
        cell = table_1.cell(3, 1)
        cell.text = slovar['admin_dep_name']

        for row in table_1.rows:
            for cell in row.cells:
                paragraph = cell.paragraphs[0]
                paragraph.runs[0].font.size = Pt(14)

        paragraph_text = f'Об увольнении {initials}'
        for paragraph in doc_1.paragraphs:
            if paragraph_text in paragraph.text:
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.bold = True
        paragraph_text = f'«Об увольнении {initials}»'
        for paragraph in doc_1.paragraphs:
            if paragraph_text in paragraph.text:
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.bold = False

        doc_1.save(f'приказы/Приказ об увольнении_{datetime.now().strftime("%Y-%m-%d %H.%M.%S")}_{initials}.docx')

    # О переводе
    def transfer(self, fio, dolzn, new_dolzn, otdel, new_otdel, data, data_zayav, slovar):
        doc_2 = docx.Document("Приказ о переводе.docx")
        for paragraph in doc_2.paragraphs:
            if "#####" in paragraph.text:
                paragraph.text = paragraph.text.replace("#####", fio)
                paragraph.runs[0].font.name = "Times New Roman"
        initials = self.format_name(fio)
        for paragraph in doc_2.paragraphs:
            if "###" in paragraph.text:
                paragraph.text = paragraph.text.replace("###", initials)
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.name = "Times New Roman"
        smen_init = self.smena(initials)
        for paragraph in doc_2.paragraphs:
            if "___" in paragraph.text:
                paragraph.text = paragraph.text.replace("___", smen_init)
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.name = "Times New Roman"
        for paragraph in doc_2.paragraphs:
            if "##" in paragraph.text:
                paragraph.text = paragraph.text.replace("##", data)
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.name = "Times New Roman"
        for paragraph in doc_2.paragraphs:
            if "-_-" in paragraph.text:
                paragraph.text = paragraph.text.replace("-_-", new_dolzn)
                paragraph.runs[0].font.name = "Times New Roman"
        for paragraph in doc_2.paragraphs:
            if "===" in paragraph.text:
                paragraph.text = paragraph.text.replace("===", new_otdel)
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.name = "Times New Roman"
        for paragraph in doc_2.paragraphs:
            if "=_=" in paragraph.text:
                paragraph.text = paragraph.text.replace("=_=", data_zayav)
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.name = "Times New Roman"
        for paragraph in doc_2.paragraphs:
            if "dolzn" in paragraph.text:
                paragraph.text = paragraph.text.replace("dolzn", dolzn)
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.name = "Times New Roman"
        for paragraph in doc_2.paragraphs:
            if "otdel" in paragraph.text:
                paragraph.text = paragraph.text.replace("otdel", otdel)
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.name = "Times New Roman"
        for paragraph in doc_2.paragraphs:
            if "head_name" in paragraph.text:
                paragraph.text = paragraph.text.replace("head_name", slovar["head_dep_name"])
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.name = "Times New Roman"
        for paragraph in doc_2.paragraphs:
            if "head_post" in paragraph.text:
                paragraph.text = paragraph.text.replace("head_post", slovar["head_dep_post"])
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.name = "Times New Roman"
        for paragraph in doc_2.paragraphs:
            if "hr_post" in paragraph.text:
                paragraph.text = paragraph.text.replace("hr_post", slovar["HR_post"])
                paragraph.runs[0].font.size = Pt(10)
                paragraph.runs[0].font.name = "Times New Roman"
        for paragraph in doc_2.paragraphs:
            if "hr_name" in paragraph.text:
                paragraph.text = paragraph.text.replace("hr_name", slovar["HR_name"])
                paragraph.runs[0].font.size = Pt(10)
                paragraph.runs[0].font.name = "Times New Roman"

        table_1 = doc_2.tables[0]
        cell = table_1.cell(1, 0)
        cell.text = slovar['Curator_post']
        cell = table_1.cell(1, 1)
        cell.text = slovar['Curator_name']

        cell = table_1.cell(2, 0)
        cell.text = slovar['Law_dep_post']
        cell = table_1.cell(2, 1)
        cell.text = slovar['Law_dep_name']

        cell = table_1.cell(3, 0)
        cell.text = slovar['admin_dep_post']
        cell = table_1.cell(3, 1)
        cell.text = slovar['admin_dep_name']

        for row in table_1.rows:
            for cell in row.cells:
                paragraph = cell.paragraphs[0]
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.name = "Times New Roman"

        paragraph_text = f'О переводе {initials}'
        for paragraph in doc_2.paragraphs:
            if paragraph_text in paragraph.text:
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.bold = True
        paragraph_text = f'«О переводе {initials}»'
        for paragraph in doc_2.paragraphs:
            if paragraph_text in paragraph.text:
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].font.bold = False

        doc_2.save(f'приказы/Приказ о переводе_{datetime.now().strftime("%Y-%m-%d %H.%M.%S")}_{initials}.docx')

    def create_word_file(self, full_name, dolzn, new_dolzn, otdel, new_otdel, data_prikaz, data_zayav, slovar, document_type):

        if document_type == "Назначение":
            self.naznach(full_name, new_dolzn, new_otdel, data_prikaz, data_zayav, slovar)
        elif document_type == "Увольнение":
            self.dismissal(full_name, dolzn, otdel, data_prikaz, data_zayav, slovar)
        elif document_type == "Перевод":
            self.transfer(full_name, dolzn, new_dolzn, otdel, new_otdel, data_prikaz, data_zayav, slovar)
        else:
            return 0


